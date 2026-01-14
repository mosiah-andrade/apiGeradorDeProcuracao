from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from io import BytesIO
from datetime import date
import os

app = FastAPI()

# --- CONFIGURAÇÃO DE CORS (Obrigatório para o React funcionar) ---
origins = ["*"] # Em produção, troque "*" pelo endereço do seu site
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# -----------------------------------------------------------------

@app.get("/health")
def health_check():
    """
    Rota leve para evitar o 'Cold Start' do Render.
    Cadastre a URL desta rota (ex: https://sua-api.onrender.com/health)
    em um serviço de monitoramento (UptimeRobot).
    """
    return {"status": "ok", "message": "API is awake and running!"}
# ==========================================

# Modelo de dados que o site vai enviar
class DadosContrato(BaseModel):
    nome: str
    cpf: str
    rg: str = "VAZIO"              # <--- NOVO
    orgao_emissor: str = "VAZIO"   # <--- NOVO
    kwp: str = "0.0"
    endereco: str = "Rua Teste"
    cidade: str = "Cidade Teste"
    classificacao: str = "B1" 
    contacontrato: str = "000000000"
    bairro: str = "Bairro Teste"
    cep: str = "00000-000"
    concessionaria: str = "TESTE"
    representante: str = "VAZIO"
    cpf_representante: str = "VAZIO"
    nome_CONTRATADO: str = "VAZIO"
    rg_CONTRATADO: str = "VAZIO"
    orgao_emissor_CONTRATADO: str = "VAZIO"
    cpf_CONTRATADO: str = "VAZIO"
    endereco_CONTRATADO: str = "VAZIO"

@app.post("/gerar-contrato/")
def gerar_contrato(dados: DadosContrato):
    # 1. Configurações Iniciais
    hoje = date.today()
    meses_pt = ['Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho', 
                'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro']
    
    # 2. Prepara os dados para substituir no Word
    mapa_dados = {
        # Dados do Cliente
        "{{NOME}}": dados.nome,
        "{{CPF}}": dados.cpf,
        "{{RG}}": dados.rg,                         
        "{{ORGAO_EMISSOR}}": dados.orgao_emissor,
        "{{ENDERECO}}": dados.endereco,
        "{{CIDADE}}": dados.cidade,
        "{{CLASSIFICACAO}}": dados.classificacao,
        "{{CONTACONTRATO}}": dados.contacontrato,
        "{{BAIRRO}}": dados.bairro,
        "{{CEP}}": dados.cep,
        "{{CONCESSIONARIA}}": dados.concessionaria,
        
        # Representante Legal (se houver)
        "{{REPRESENTANTE}}": dados.representante,
        "{{CPF_DO_REPRESENTANTE}}": dados.cpf_representante,

        # --- NOVOS CAMPOS DO CONTRATADO ---
        "{{NOME_CONTRATADO}}": dados.nome_CONTRATADO,
        "{{RG_CONTRATADO}}": dados.rg_CONTRATADO,
        "{{ORGAO_EMISSOR_CONTRATADO}}": dados.orgao_emissor_CONTRATADO,
        "{{CPF_CONTRATADO}}": dados.cpf_CONTRATADO,
        "{{ENDERECO_CONTRATADO}}": dados.endereco_CONTRATADO,
        
        # Datas
        "{{DIA}}": str(hoje.day),
        "{{MES}}": meses_pt[hoje.month - 1],  
        "{{ANO}}": str(hoje.year),
    }

    # 3. Lógica de Seleção do Modelo (CPF vs CNPJ)
    # Remove pontos, traços e barras para contar apenas os números
    doc_limpo = dados.cpf.replace(".", "").replace("-", "").replace("/", "").strip()

    if len(doc_limpo) > 11:
        # Se tiver mais de 11 dígitos, é CNPJ (Pessoa Jurídica)
        caminho_modelo = "modelo_pj.docx" 
        print(f"Detectado CNPJ ({doc_limpo}). Usando modelo PJ.")
    else:
        # Caso contrário, assumimos que é CPF (Pessoa Física)
        caminho_modelo = "modelo_pf.docx"
        print(f"Detectado CPF ({doc_limpo}). Usando modelo PF.")
    
    # Verifica se o arquivo existe antes de tentar abrir
    if not os.path.exists(caminho_modelo):
        raise HTTPException(
            status_code=404, 
            detail=f"ERRO: O arquivo '{caminho_modelo}' não foi encontrado na pasta da API. Verifique se ele existe!"
        )

    # --- 4. Abre e substitui (COM FORMATAÇÃO) ---
    doc = Document(caminho_modelo)

    # Função auxiliar para substituir mantendo estilo
    def substituir_nas_runs(container, mapa):
        for run in container.runs:
            for codigo, valor in mapa.items():
                if codigo in run.text:
                    run.text = run.text.replace(codigo, str(valor))

    # Aplica nos parágrafos
    for paragrafo in doc.paragraphs:
        substituir_nas_runs(paragrafo, mapa_dados)
    
    # Aplica nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                # Células também têm parágrafos dentro
                for paragrafo in celula.paragraphs:
                    substituir_nas_runs(paragrafo, mapa_dados)

    # 5. Salva na memória e envia
    arquivo_memoria = BytesIO()
    doc.save(arquivo_memoria)
    arquivo_memoria.seek(0)

    nome_arquivo = f"Procuracao_{dados.nome.replace(' ', '_')}.docx"
    
    return StreamingResponse(
        arquivo_memoria, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={nome_arquivo}"}
    )